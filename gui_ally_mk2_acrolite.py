from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl.styles import Font
from openpyxl import Workbook
from PyQt5 import QtWidgets
from docx import Document
import docx2txt
import csv
import copy
import time

def printme(docs_scanned):
    print("SEE ME")
    print("DOCS SCANNED IS " + str(docs_scanned))

def initial_labels(win):
    plug = QtWidgets.QLabel(win)  # make a label in window
    plug.setText("Product of SpaceCows™ est. 2020")  # label is named this
    plug.setGeometry(160,10,300,30)
    version = QtWidgets.QLabel(win)
    version.setText("Acrolite Version 2.1")
    version.setGeometry(230,50,300,30)

def define_file_path():
    breakout = False
    while not breakout:
        name = str(input("\nPlease enter your file name, including the suffix (ie Word.docx): "))
        if "." in name:
            breakout=True
        else:
            print("No suffix detected. Be sure to include the suffix in the name.")

    doc2bscanned = "hi there"
    breakout = False
    while not breakout:
        path = str(input("\nPlease provide the absolute path to the document you'd like to scan. DO NOT include the actual file name at the far right.\n(Press 'H' for help)\nPATH: "))
        if path == "H":
            print("To get the path of a file, simply right click it and go to 'properties' at the bottom. Click it.")
            print("Once in properties, look for the 'Location' information. It should be in the form 'C:\\Users\\e##...")
            print("Copy the part, and paste it into the prompt below.")
        else:
            doc2bscanned = r""+str(path)+"\\"+str(name)
            breakout = True

    # print("SCAN THIS DOC " +str(doc2bscanned))

    return doc2bscanned

def collect_doc_words(doc2bscanned):
    start = time.time()
    result = docx2txt.process(doc2bscanned)
    finish = time.time()

    counted_words = count_doc_words(doc2bscanned)
    counting_time = finish-start

    scan_performance = str(counted_words)+" words scanned in " + str(counting_time) + " seconds"


    print(scan_performance)
    words = result.split()

    return words,scan_performance

def detect_pairs_above(words):  # detecting definitions in the actual document
    definitions = []
    for word in words:
        local_definition = []
        left = 0
        right = 0
        uppers = 0
        lowers = 0
        nums = 0
        for letter in word:
            if letter.isupper():
                uppers += 1
            elif letter.islower():
                lowers += 1
            elif letter.isnumeric():
                nums +=1
        if "(" in word:
            left = 1
        if ")" in word:
            right = 1
        if left == 1 and right == 1 and uppers >= 2 and lowers <=2 and nums <= 2:  # we have determined this word is an acronym
            curr_index = words.index(word)
            if curr_index > 0:
                breakout = False
                back = 1
                while not breakout: # criteria for stopping looking for a definition
                    if curr_index - back == 0:
                        breakout = True
                    elif words[curr_index - back].islower():
                        if words[curr_index-back] == "and":
                            pass
                        elif words[curr_index-back] == "of":
                            pass
                        else:
                            breakout = True
                    elif words[curr_index - back].isupper():
                        breakout = True
                    elif words[curr_index - back] == ")":
                        breakout = True
                    elif words[curr_index - back] == ",":
                        breakout = True
                    elif words[curr_index - back].isnumeric():
                        breakout = True
                    elif "." in words[curr_index - back]:
                        breakout = True
                    local_definition.append(words[curr_index - back])
                    back += 1

            # delete last thing appended, always appended unintentionally
            del local_definition[len(local_definition) - 1]

            # clean acronym
            word = ''.join(e for e in word if e.isalnum())

            # clean definition
            for i in range(len(local_definition) - 1, -1, -1):
                if local_definition[i] in ["and", "&", ".", "-", " ","The","/"] and i == len(local_definition) - 1:
                    del local_definition[i]
                elif local_definition[i] in ["and", "&", ".", "-", " ","The","/"] and i == 0:
                    del local_definition[i]

            local_definition.append(word)
            local_definition.reverse()
            definitions.append(local_definition)

    return definitions

def data_processing(definitions):

    # convert definition's lists to a string
    definitions = def_list_to_string(definitions)
    definitions = no_dupes(definitions)

    # compare blacklist with new acronyms
    blacklist = csv_to_blacklist()
    non_blacklists = []
    for definition in definitions:
        if definition[0] in blacklist:
            # print(str(definition[0]) + " is blacklisted.")
            pass
        else:
            non_blacklists.append(definition)
    definitions = copy.deepcopy(non_blacklists)

    # establish glossary values
    dictionary = read_from_csv(dictionary={}, csv_doc="lite_gloss.csv")
    glossary_values = dict_to_list(dictionary)

    # determine what's in the glossary and what's not
    gloss_matches = list_overlap(glossary_values, definitions)

    for match in gloss_matches:
        definitions.remove(match)

    follow_ups = copy.deepcopy(definitions)

    return gloss_matches,follow_ups

def user_affirmation(gloss_matches,follow_ups):

    # initialize customer and glossary dictionary
    to_customer = list_to_dict(gloss_matches)
    to_glossary = list_to_dict(gloss_matches)

    # initialize blacklist for appending
    blacklist = csv_to_blacklist()

    # initiate user conversation
    for follow_up in follow_ups:
        breakout = False
        if follow_up[1] == "":
            while not breakout:
                try:
                    print("\n\n--No definition for '" + str(follow_up[0]) + "' was found--")
                    definition = str(
                        input("What is the definition?\n[B]lacklist\n[Enter]Skip\nOR Enter the definition: \n"))
                    if definition in ["B", "b", "bl", "Bl", "BL", "Black", "BLACK", "black"]:
                        print("Blacklisting " + str(follow_up[0]) + "...")
                        blacklist.append(follow_up[0])
                        breakout = True
                    elif definition == "":
                        print("Skipping " + str(follow_up[0]) + "...")
                        breakout = True
                    else:
                        print("Definition for " + str(follow_up[0]) + " is '" + str(definition)+"'")
                        to_customer[follow_up[0]]=definition
                        to_glossary[follow_up[0]]=definition
                        breakout = True
                except Exception:
                    print("\n/PLEASE ENTER A VALID INPUT./")

        else:
            while not breakout:
                try:
                    pairing = str(input("\n\nFound pair ['" + str(follow_up[0]) + "' : '"+str(follow_up[1])+"']\nIs this pairing correct?\n[Y]es\n[D]efinition Wrong\n[A]cronym Nonexistent\n[Enter]Skip\n"))
                    if pairing in ["y","Y","yes","YES","Yes"]:
                        print("Sent to Customer List")
                        to_customer[follow_up[0]] = follow_up[1]
                        to_glossary[follow_up[0]] = follow_up[1]
                        breakout = True
                    elif pairing in ["D","d","def","Def","definition","Definition","DEF","DEFINITION"]:
                        # print("Wrong definition")
                        correction = str(input("->Can you give the right definition?\n->[Enter]Skip\n->Definition: "))
                        if correction == "":
                            print("Skipped '[" + str(follow_up[0]) + "' : '"+str(follow_up[1])+"]'")
                        else:
                            to_customer[follow_up[0]]=correction
                            to_glossary[follow_up[0]] = correction
                            print("Sent to Customer List")
                        breakout = True
                    elif pairing in ["A","a","acro","Acro","ACRO","Acronym","acronym","ACRONYM"]:
                        print("Blacklisting...")
                        blacklist.append(follow_up[0])
                        breakout = True
                    elif pairing == "":
                        print("Skipped '[" + str(follow_up[0]) + "' : '"+str(follow_up[1])+"]'")
                        breakout = True
                    else:
                        raise Exception
                except Exception:
                    print("\n/PLEASE ENTER A VALID INPUT./")

    # send blacklist back
    blacklist_to_csv(blacklist)

    print("TO GLOSS IS " + str(to_glossary))

    write_to_csv(to_glossary,"lite_gloss.csv")

    return to_customer

def glossary_definition_overlay_splitter(glossary,definitions):
    gloss_matches = []
    follow_ups = []

    for term in glossary:
        for definition in definitions:
            if term == definition:
                gloss_matches.append(term)
            else:
                follow_ups.append(definition)
    return gloss_matches,follow_ups

def list_overlap(list1, list2):
    lst3 = [value for value in list1 if value in list2]
    return lst3

def list_to_dict(list):
    dictionary = {}
    for item in list:
        dictionary[item[0]]=item[1:][0]
    return dictionary

def dict_to_list(dictionary):
    large_pairs = []
    for key in dictionary:
        small_pairs = []
        small_pairs.append(key)
        small_pairs.append(dictionary[key])
        large_pairs.append(small_pairs)
    return large_pairs

def def_list_to_string(definitions):
    for definition in definitions:
        defy = " ".join(definition[1:])
        del definition[1:]
        definition.append(defy)
    return definitions

def no_dupes(definitions): # use for finding similairites in glossary and doc
    no_dupes = [x for n, x in enumerate(definitions) if x not in definitions[:n]]
    return no_dupes

def dupes(definitions): # use for finding similairites in glossary and doc
    dupes = [x for n, x in enumerate(definitions) if x in definitions[:n]]
    return dupes

def write_to_csv(dictionary, csv_doc):
    with open(csv_doc, 'w') as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(["Term", "Description"])
        for key, value in dictionary.items():
            writer.writerow([key, value])

def csv_to_blacklist():
    with open('blacklistlite.csv', newline='') as f:
        reader = csv.reader(f)
        blacklist = next(reader)
    return blacklist

def blacklist_to_csv(blacklist):
    with open('blacklistlite.csv', 'w') as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(blacklist)
    return

def read_from_csv(dictionary,csv_doc):
    with open(csv_doc, newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            dictionary[row['Term']] = row['Description']
    return dictionary

def sort_database_alphabetically(acronyms):
    sorted_list = {}
    for key, value in sorted(acronyms.items()):
        sorted_list[key] = value
    acronyms = {}
    acronyms = copy.deepcopy(sorted_list)
    return acronyms

def count_doc_words(document):
    docu = open(document, "rb")
    document = Document(docu)

    count = 0
    for para in document.paragraphs:
        content= para.text
        countage= len(content.split())
        count = count + countage
    return count

def glossary_to_excel(dictionary):
    workbook = Workbook()
    sheet = workbook.active

    sheet["A1"] = "Acronym"
    sheet["B1"] = "Definition"

    bold_font = Font(bold=True)
    sheet["A1"].font = bold_font
    sheet["B1"].font = bold_font

    count = 2
    for key, value in dictionary.items():
        sheet["A"+str(count)] = str(key)
        sheet["B" + str(count)] = str(value)
        count += 1

    workbook.save(filename="Customer_List.xlsx")

def glossary_to_word(acronym_glossary):

    print("Glossary is " + str(acronym_glossary))

    # put in alphabetical order
    acronym_glossary = sort_database_alphabetically(acronym_glossary)

    # initialize document
    doc = Document()
    table = doc.add_table(rows=0,cols=2)

    # add titles
    row = table.add_row().cells
    keys = row[0].add_paragraph("")
    keys.add_run('Acronyms').bold = True
    keys.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    defs = row[1].add_paragraph("")
    defs.add_run('Definitions').bold = True
    defs.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # add dictionary
    for key in acronym_glossary:

        row = table.add_row().cells
        keys = row[0].add_paragraph(key)
        keys.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        defs = row[1].add_paragraph(acronym_glossary[key])
        defs.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # save doc
    doc.save('Customer_List.docx')

